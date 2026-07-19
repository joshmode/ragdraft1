import io
import os
import re
import tempfile


def generate_docx(md_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not md_text or len(md_text.strip()) < 20:
        raise ValueError("Generated document is empty.")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    def add_hr():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(3)
        ppr = paragraph._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        border.append(bottom)
        ppr.append(border)

    def add_runs(paragraph, raw_line: str, size: float = 10.5, bold: bool = False):
        for part in re.split(r"(\*\*.*?\*\*|\*.*?\*|_.*?_)", raw_line):
            if part.startswith("**") and part.endswith("**"):
                text = part[2:-2]
                is_bold = True
                is_italic = False
            elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
                text = part[1:-1]
                is_bold = bold
                is_italic = True
            else:
                text = part
                is_bold = bold
                is_italic = False
            run = paragraph.add_run(text)
            run.bold = is_bold
            run.italic = is_italic
            run.font.size = Pt(size)

    for line in md_text.split("\n"):
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x0D, 0x0F, 0x11)
        elif line.startswith("## "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            run = paragraph.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(10)
            add_hr()
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            add_runs(paragraph, line[4:].strip(), bold=True)
        elif re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", line.strip()):
            add_hr()
        else:
            match = re.match(r"^(\s*)[-*•]\s+(.*)", line)
            paragraph = doc.add_paragraph(style="List Bullet") if match else doc.add_paragraph()
            if match:
                paragraph.paragraph_format.left_indent = Inches(0.2)
                add_runs(paragraph, match.group(2).strip())
            else:
                add_runs(paragraph, line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf(md_text: str) -> bytes:
    from markdown_pdf import MarkdownPdf
    try:
        from markdown_pdf import Section
    except ImportError:
        from markdown_pdf.Section import Section

    pdf = MarkdownPdf(toc_level=0)
    pdf.add_section(Section(md_text))
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        path = tmp.name
    try:
        pdf.save(path)
        with open(path, "rb") as file:
            return file.read()
    finally:
        if os.path.exists(path):
            os.remove(path)
