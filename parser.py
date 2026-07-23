# pyrefly
import fitz
import importlib.util
import re
import csv
import io
import zipfile
from dataclasses import dataclass, field

# find_spec checks availability without importing torch (700MB+ RSS) until a scan actually needs OCR
def _ocr_installed() -> bool:
    return (
        importlib.util.find_spec("easyocr") is not None
        and importlib.util.find_spec("pdf2image") is not None
        and importlib.util.find_spec("numpy") is not None
    )


# easyocr Reader is expensive to init, keep it as a per-process singleton once loaded
_reader = None

def _get_reader():
    global _reader
    if _reader is None:
        import easyocr
        _reader = easyocr.Reader(["en"], gpu=False)  # gpu=True if you have cuda gpu eg nvidia rtx serise
    return _reader


EMAIL_RE    = re.compile(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}')
PHONE_RE    = re.compile(r'(\+?\d[\d\s\-().]{7,}\d)')
_YEAR_RANGE_RE = re.compile(r'^\(?(19|20)\d{2}\)?\s*[-–—]\s*\(?(19|20)\d{2}\)?$')
LINKEDIN_RE = re.compile(r'linkedin\.com/in/[\w\-]+', re.IGNORECASE)
GITHUB_RE   = re.compile(r'github\.com/[\w\-]+', re.IGNORECASE)
URL_RE      = re.compile(r'https?://\S+')

# maps header permutations to section names 
SECTION_ALIASES = {
    "experience":              "EXPERIENCE",
    "work experience":         "EXPERIENCE",
    "professional experience": "EXPERIENCE",
    "employment history":      "EXPERIENCE",
    "work history":            "EXPERIENCE",
    "internships":             "EXPERIENCE",
    "internship experience":   "EXPERIENCE",
    "relevant experience":     "EXPERIENCE",
    "career history":          "EXPERIENCE",

    "education":               "EDUCATION",
    "academic background":     "EDUCATION",
    "qualifications":          "EDUCATION",
    "academic qualifications": "EDUCATION",

    "skills":                  "SKILLS",
    "technical skills":        "SKILLS",
    "core competencies":       "SKILLS",
    "technologies":            "SKILLS",
    "tools & technologies":    "SKILLS",
    "tools and technologies":  "SKILLS",
    "key skills":              "SKILLS",
    "competencies":            "SKILLS",
    "technical competencies":  "SKILLS",
    "areas of expertise":      "SKILLS",
    "skill set":               "SKILLS",

    "projects":                "PROJECTS",
    "personal projects":       "PROJECTS",
    "side projects":           "PROJECTS",
    "portfolio":               "PROJECTS",
    "key projects":            "PROJECTS",
    "academic projects":       "PROJECTS",
    "relevant projects":       "PROJECTS",
    "selected projects":       "PROJECTS",
    "project experience":      "PROJECTS",
    "notable projects":        "PROJECTS",

    "summary":                 "SUMMARY",
    "professional summary":    "SUMMARY",
    "profile":                 "SUMMARY",
    "objective":               "SUMMARY",
    "about me":                "SUMMARY",
    "career objective":        "SUMMARY",
    "career summary":          "SUMMARY",
    "executive summary":       "SUMMARY",
    "personal statement":      "SUMMARY",

    "certifications":          "CERTIFICATIONS",
    "certificates":            "CERTIFICATIONS",
    "licenses":                "CERTIFICATIONS",
    "licenses & certifications": "CERTIFICATIONS",
    "professional certifications": "CERTIFICATIONS",

    "awards":                  "AWARDS",
    "honors":                  "AWARDS",
    "achievements":            "AWARDS",
    "honors & awards":         "AWARDS",
    "awards & recognition":    "AWARDS",

    "publications":            "PUBLICATIONS",
    "research":                "PUBLICATIONS",
    "research experience":     "PUBLICATIONS",
    "papers":                  "PUBLICATIONS",

    "volunteer":               "VOLUNTEER",
    "volunteering":            "VOLUNTEER",
    "community":               "VOLUNTEER",
    "community service":       "VOLUNTEER",
    "volunteer experience":    "VOLUNTEER",
    "community involvement":   "VOLUNTEER",

    "languages":               "LANGUAGES",
    "interests":               "INTERESTS",
    "hobbies":                 "INTERESTS",
    "activities":              "INTERESTS",
    "extracurricular":         "INTERESTS",
    "extracurricular activities": "INTERESTS",
    "references":              "REFERENCES",
}

# remove markdown formatting issues from pdf text extracted
_MD_STRIP_RE = re.compile(r'[\*_#`]+')
_DECO_RE = re.compile(r'^[\s\u200b\ufeff\u00a0│|▌►▶◆■□●○]+|[\s\u200b\ufeff\u00a0│|▌►▶◆■□●○]+$')

_SKILL_SUBLABEL_RE = re.compile(
    r'^\s*(?:languages|frameworks|libraries|databases|tools|platforms|'
    r'ai|ml|devops|cloud|operating\s*systems|software|hardware|'
    r'technologies|methods|methodologies|interests|publications|'
    r'research|areas|certifications|awards|honors|activities|'
    r'frameworks\s*[&]\s*libraries|tools\s*[&]\s*(?:methods|technologies))'
    r'\s*[:/|•·–—-]\s*\S',
    re.IGNORECASE,
)

_SUBSECTION_OVERRIDES = {
    "LANGUAGES", "INTERESTS", "PUBLICATIONS", "CERTIFICATIONS",
    "AWARDS", "VOLUNTEER", "REFERENCES",
}

_EMPTY_BULLET_RE = re.compile(r'^\s*[•‣▪▫◦●○\-*]\s*[\u200b\ufeff]*\s*$')


_LETTER_SPACED_RE = re.compile(r'^(?:[A-Za-z](?:\s+|$)){2,}$')


def _declump_letters(text: str) -> str:
    """collapse letter-spaced/tracked headers like 'T E C H N I C A L' into 'TECHNICAL'."""
    words = []
    cur = []
    for token in text.split(' '):
        if len(token) == 1 and token.isalpha():
            cur.append(token)
        else:
            if cur:
                words.append(''.join(cur))
                cur = []
            if token:
                words.append(token)
    if cur:
        words.append(''.join(cur))
    return ' '.join(words)


def _clean_header(line: str) -> str:
    """strip formatting artifacts for section header matching."""
    out = _MD_STRIP_RE.sub('', line)
    out = _DECO_RE.sub('', out).strip()
    out = re.sub(r'[\s:;\-–—|]+$', '', out)
    if _LETTER_SPACED_RE.match(out) and len(out.replace(' ', '')) >= 4:
        out = _declump_letters(out)
    return out


# regex built from aliases
_SECTION_RE = re.compile(
    r'^(' + '|'.join(re.escape(k) for k in sorted(SECTION_ALIASES.keys(), key=len, reverse=True)) + r')\b.*$',
    re.IGNORECASE,
)

# single-word canonical starters
_CANON_STARTERS = {
    word.split()[0].lower()
    for word in SECTION_ALIASES.keys()
    if len(word.split()) == 1
}
_CANON_STARTERS.update({
    "work", "professional", "employment", "career",
    "technical", "core", "key", "academic", "personal",
    "side", "selected", "notable", "relevant",
    "volunteer", "community", "extracurricular",
})


@dataclass
class ParsedResume:
    raw_text: str = ""
    contact:  dict = field(default_factory=dict)
    sections: dict[str, list[str]] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    ocr_used: bool = False

    def section_text(self, name: str) -> str:
        return "\n".join(self.sections.get(name.upper(), []))

    def to_llm_prompt(self) -> str:
        parts = []
        if self.contact:
            parts.append("=== CONTACT ===")
            for key, val in self.contact.items():
                parts.append(f"{key}: {val}")
        for sec, lines in self.sections.items():
            parts.append(f"\n=== {sec} ===")
            parts.extend(lines)
        return "\n".join(parts)


def _is_scan(lines: list[str]) -> bool:
    useful = [l for l in lines if len(l.strip()) > 3]
    return len(useful) < 10

def _ocr_fallback(raw: bytes) -> list[str]:
    if not _ocr_installed():
        return []

    import numpy as np
    import pdf2image

    reader = _get_reader()
    images = pdf2image.convert_from_bytes(raw, dpi=300)
    all_lines = []
    for img in images:
        arr = np.array(img)
        # detail=1 gives confidence scores, sort by y for reading order
        results = reader.readtext(arr, detail=1, paragraph=False)
        results.sort(key=lambda r: r[0][0][1])

        for (bbox, text, conf) in results:
            cleaned = text.strip()
            if conf > 0.3 and len(cleaned) > 2:
                all_lines.append(cleaned)

    return all_lines


def _sorted_blocks(page: fitz.Page) -> list[str]:
    blocks = page.get_text("blocks")
    txt_blocks = [b for b in blocks if b[6] == 0]
    txt_blocks.sort(key=lambda b: (round(b[1] / 10), b[0]))

    lines = []
    for blk in txt_blocks:
        for raw in blk[4].split("\n"):
            s = raw.strip()
            if s and not _EMPTY_BULLET_RE.match(s):
                lines.append(s)
    return lines


def _match_section(line: str) -> str | None:
    cleaned = _clean_header(line)
    if not cleaned:
        return None

    words = cleaned.split()

    if len(words) <= 8:
        m = _SECTION_RE.match(cleaned)
        if m:
            return SECTION_ALIASES[m.group(1).lower()]

    lower = cleaned.lower().strip()
    if lower in SECTION_ALIASES:
        return SECTION_ALIASES[lower]

    stripped = re.sub(r'[\s:;\-–—|]+$', '', lower).strip()
    stripped = re.sub(r'\s*\d+\s*$', '', stripped).strip()
    if stripped in SECTION_ALIASES:
        return SECTION_ALIASES[stripped]

    # cleanse colon that may be embedded
    colon_stripped = re.sub(r'\s*:\s*.*$', '', lower).strip()
    if colon_stripped in SECTION_ALIASES:
        return SECTION_ALIASES[colon_stripped]

    if 1 <= len(words) <= 8:
        first = words[0].lower().rstrip(':')
        stylized = (
            cleaned.isupper()
            or cleaned.istitle()
            or (len(words) <= 5 and all(w[0].isupper() for w in words if w and w[0].isalpha()))
            or cleaned.replace(' ', '').isupper()
        )
        if first in _CANON_STARTERS and stylized:
            for n in range(len(words), 0, -1):
                cand = ' '.join(w.lower().rstrip(':') for w in words[:n])
                if cand in SECTION_ALIASES:
                    return SECTION_ALIASES[cand]

        for n in range(min(len(words), 4), 0, -1):
            cand = ' '.join(w.lower().rstrip(':') for w in words[:n])
            if cand in SECTION_ALIASES:
                rest = words[n:]
                if not rest or all(not re.search(r'[.!?]', w) for w in rest):
                    return SECTION_ALIASES[cand]

        # try matching each word to aliase
        for n in range(len(words), 0, -1):
            cand = ' '.join(w.lower().strip(':.;-–—|') for w in words[:n])
            if cand in SECTION_ALIASES:
                return SECTION_ALIASES[cand]

        if 1 <= len(words) <= 6:
            last = words[-1].lower().rstrip(':')
            stylized = (
                cleaned.istitle()
                or cleaned.isupper()
                or all(w[0].isupper() for w in words if w and w[0].isalpha())
            )
            if last in SECTION_ALIASES and stylized:
                return SECTION_ALIASES[last]

        if len(words) == 1 and len(words[0]) >= 8:
            glued = words[0].rstrip(':')
            split = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', glued)
            if ' ' in split:
                return _match_section(split)
            glued_lower = glued.lower()
            for alias, canon in SECTION_ALIASES.items():
                if ' ' in alias and alias.replace(' ', '') == glued_lower:
                    return canon

        return None


def _is_header_fragment(line: str) -> bool:
    """true if line looks like the first half of a multiword header wrapped onto two sep lines."""
    cleaned = _clean_header(line)
    words = cleaned.split()
    if len(words) != 1:
        return False
    word = words[0].rstrip(':')
    if not word.isalpha():
        return False
    stylized = cleaned.isupper() or cleaned.istitle()
    return stylized and word.lower() in _CANON_STARTERS


def _parse_contact(lines: list[str]) -> dict:
    contact = {}
    for i, line in enumerate(lines[:15]):
        if i == 0 and not EMAIL_RE.search(line) and not PHONE_RE.search(line):
            contact["name"] = line

        if "email" not in contact:
            m = EMAIL_RE.search(line)
            if m:
                contact["email"] = m.group()

        if "phone" not in contact:
            m = PHONE_RE.search(line)
            if m:
                candidate = m.group().strip()
                digits = re.sub(r'\D', '', candidate)
                if len(digits) >= 7 and not _YEAR_RANGE_RE.match(candidate):
                    contact["phone"] = candidate

        if "linkedin" not in contact:
            m = LINKEDIN_RE.search(line)
            if m:
                contact["linkedin"] = m.group()

        if "github" not in contact:
            m = GITHUB_RE.search(line)
            if m:
                contact["github"] = m.group()

        if "website" not in contact:
            m = URL_RE.search(line)
            if m and "linkedin" not in m.group() and "github" not in m.group():
                contact["website"] = m.group()

    return contact


def _lines_to_resume(all_lines: list[str], result: ParsedResume | None = None) -> ParsedResume:
    if result is None:
        result = ParsedResume()

    result.raw_text = "\n".join(all_lines)
    result.contact  = _parse_contact(all_lines)

    sections: dict[str, list[str]] = {"HEADER": []}
    current = "HEADER"

    skip_next = False
    for idx, line in enumerate(all_lines):
        if skip_next:
            skip_next = False
            continue

        canonical = _match_section(line)
        if not canonical and idx + 1 < len(all_lines) and _is_header_fragment(line):
            nxt = all_lines[idx + 1]
            merged = f"{_clean_header(line)} {_clean_header(nxt)}".strip()
            canonical = _match_section(merged)
            if canonical:
                skip_next = True

        if canonical and canonical != current:
            if current in ("SKILLS", "EDUCATION") and canonical in _SUBSECTION_OVERRIDES:
                if _SKILL_SUBLABEL_RE.match(line):
                    sections[current].append(line)
                    continue
            current = canonical
            if current not in sections:
                sections[current] = []
        else:
            # canonical == current (e.g. a "Technical Skills:" sub-line inside a SKILLS
            # section, since that alias also maps to SKILLS) is a content line for the
            # section we're already in, not a new section - keep it, don't drop it.
            sections[current].append(line)

    contact_vals = set(result.contact.values())
    sections["HEADER"] = [
        l for l in sections.get("HEADER", [])
        if l not in contact_vals
        and not EMAIL_RE.search(l)
        and not PHONE_RE.search(l)
        and not LINKEDIN_RE.search(l)
        and not GITHUB_RE.search(l)
    ]

    result.sections = {k: v for k, v in sections.items() if v}

    for must_have in ("EXPERIENCE", "EDUCATION", "SKILLS"):
        if must_have not in result.sections:
            result.warnings.append(f"'{must_have}' section not detected: the header might not be interpretable, please check your resume accordingly.")

    return result


def parse_pdf(pdf_file) -> ParsedResume:
    result = ParsedResume()
    raw = pdf_file.read() if hasattr(pdf_file, "read") else pdf_file

    try:
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as e:
        result.warnings.append(f"couldn't open PDF: {e}")
        return result

    all_lines: list[str] = []
    try:
        for page in doc:
            all_lines.extend(_sorted_blocks(page))
    finally:
        doc.close()

    if _is_scan(all_lines):
        if not _ocr_installed():
            result.warnings.append("PDF appears scanned but easyocr isn't installed. \n Run: pip install easyocr pdf2image Pillow numpy \n Also install poppler for pdf2image, open README for instructions")
            return result

        result.warnings.append("No readable text found, running OCR. Please wait as the module loads.")
        all_lines = _ocr_fallback(raw)
        result.ocr_used = True

        if not all_lines:
            result.warnings.append("OCR couldn't parse text. PDF may be corrupt or excessively styled.")
            return result

    return _lines_to_resume(all_lines, result)


def parse_docx(docx_file) -> ParsedResume:
    try:
        from docx import Document
    except ImportError:
        result = ParsedResume()
        result.warnings.append("python-docx not installed. run: pip install python-docx")
        return result

    raw = docx_file.read() if hasattr(docx_file, "read") else docx_file
    import io
    try:
        doc = Document(io.BytesIO(raw) if isinstance(raw, bytes) else raw)

        all_lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_lines.append(text)

        return _lines_to_resume(all_lines)
    except Exception as e:
        result = ParsedResume()
        result.warnings.append(f"couldn't open DOCX: {e}")
        return result


def parse_text(txt_file) -> ParsedResume:
    raw = txt_file.read() if hasattr(txt_file, "read") else txt_file
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")

    all_lines = [line.strip() for line in raw.splitlines() if line.strip()]
    return _lines_to_resume(all_lines)


def parse_markdown(md_file) -> ParsedResume:
    raw = md_file.read() if hasattr(md_file, "read") else md_file
    if isinstance(raw, bytes):
        raw = raw.decode("utf-8", errors="replace")

    all_lines: list[str] = []
    for line in raw.splitlines():
        cleaned = re.sub(r'^#+\s+', '', line)
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)
        cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)
        cleaned = re.sub(r'_(.+?)_', r'\1', cleaned)
        cleaned = re.sub(r'`(.+?)`', r'\1', cleaned)
        cleaned = re.sub(r'>\s*', '', cleaned)
        cleaned = cleaned.strip()
        if cleaned:
            all_lines.append(cleaned)

    return _lines_to_resume(all_lines)


def parse_odt(odt_file) -> ParsedResume:
    try:
        from odf.opendocument import load
        from odf import text as odf_text
        from odf import teletype
    except ImportError:
        result = ParsedResume()
        result.warnings.append("odfpy not installed. run: pip install odfpy")
        return result

    raw = odt_file.read() if hasattr(odt_file, "read") else odt_file
    try:
        import io
        doc = load(io.BytesIO(raw) if isinstance(raw, bytes) else raw)
        nodes = doc.getElementsByType(odf_text.H) + doc.getElementsByType(odf_text.P)
        all_lines = [teletype.extractText(node).strip() for node in nodes]
        all_lines = [line for line in all_lines if line]
        return _lines_to_resume(all_lines)
    except Exception as e:
        result = ParsedResume()
        result.warnings.append(f"couldn't open ODT: {e}")
        return result


def parse_legacy_doc(_doc_file) -> ParsedResume:
    result = ParsedResume()
    result.warnings.append("Legacy .doc files are not supported. Please save the file as .docx, .pdf, .txt, or .odt.")
    return result


def parse_linkedin_export(export_file) -> ParsedResume:
    raw = export_file.read() if hasattr(export_file, "read") else export_file
    result = ParsedResume()
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            entries = [entry for entry in archive.infolist() if not entry.is_dir()]
            if len(entries) > 100 or sum(entry.file_size for entry in entries) > 25 * 1024 * 1024:
                result.warnings.append("LinkedIn export is too large to process.")
                return result

            files = {
                entry.filename.rsplit("/", 1)[-1].lower(): entry
                for entry in entries
                if entry.filename.lower().endswith(".csv")
            }

            def rows(name: str) -> list[dict]:
                entry = files.get(name)
                if not entry:
                    return []
                text = archive.read(entry).decode("utf-8-sig", errors="replace")
                return list(csv.DictReader(io.StringIO(text)))

            profile = rows("profile.csv")
            positions = rows("positions.csv")
            education = rows("education.csv")
            skills = rows("skills.csv")

            if profile:
                first = profile[0]
                name = " ".join(part for part in [first.get("First Name", "").strip(), first.get("Last Name", "").strip()] if part)
                if name:
                    result.contact["name"] = name

            sections: dict[str, list[str]] = {}
            if positions:
                lines = []
                for item in positions:
                    title = item.get("Title", "").strip()
                    company = item.get("Company Name", "").strip()
                    started = item.get("Started On", "").strip()
                    finished = item.get("Finished On", "").strip()
                    heading = " | ".join(part for part in [title, company, " - ".join(part for part in [started, finished] if part)] if part)
                    if heading:
                        lines.append(heading)
                    description = item.get("Description", "").strip()
                    if description:
                        lines.extend(line.strip() for line in description.splitlines() if line.strip())
                if lines:
                    sections["EXPERIENCE"] = lines

            if education:
                lines = []
                for item in education:
                    text = " | ".join(part for part in [item.get("Degree Name", "").strip(), item.get("School Name", "").strip(), item.get("Notes", "").strip()] if part)
                    if text:
                        lines.append(text)
                if lines:
                    sections["EDUCATION"] = lines

            if skills:
                lines = [item.get("Name", "").strip() for item in skills]
                lines = [line for line in lines if line]
                if lines:
                    sections["SKILLS"] = [", ".join(lines)]

            if not sections:
                result.warnings.append("No resume sections were found in this LinkedIn export.")
                return result

            result.sections = sections
            parts = []
            if result.contact.get("name"):
                parts.append(result.contact["name"])
            for section, lines in sections.items():
                parts.append(section)
                parts.extend(lines)
            result.raw_text = "\n".join(parts)
            return result
    except zipfile.BadZipFile:
        result.warnings.append("LinkedIn exports must be a valid ZIP archive.")
        return result
    except Exception as e:
        result.warnings.append(f"Couldn't parse LinkedIn export: {e}")
        return result


_FORMAT_PARSERS = {
    ".pdf":  parse_pdf,
    ".docx": parse_docx,
    ".doc":  parse_legacy_doc,
    ".odt":  parse_odt,
    ".txt":  parse_text,
    ".md":   parse_markdown,
    ".zip":  parse_linkedin_export,
}


def parse_file(file_data, filename: str = "") -> ParsedResume:
    import os as _os
    ext = _os.path.splitext(filename)[1].lower() if filename else ".pdf"
    parser_fn = _FORMAT_PARSERS.get(ext)
    if parser_fn is None:
        result = ParsedResume()
        result.warnings.append(f"Unsupported resume format: {ext or 'unknown'}.")
        return result
    return parser_fn(file_data)
