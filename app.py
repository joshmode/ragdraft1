import sys
import html
import base64
import io
import json
import os
from typing import Any
from dataclasses import dataclass, field

try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ImportError:
    STREAMLIT_AVAILABLE = False
    print("streamlit is required. run: pip install streamlit")
    sys.exit(1)

try:
    from parser import parse_file
    from analyser import analyse, gen_cv, gen_cover_letter
    import auth
    import db
    import mentor
    import job_scraper
    import feedback
    from router import _is_key_placeholder
    BACKEND_AVAILABLE = True
except Exception as e:
    print(f"backend import error: {e}")
    BACKEND_AVAILABLE = False

_PAGE_CONFIG: dict[str, str] = {
    "page_title":            "RAGsToRiches",
    "page_icon":             "RT",
    "layout":                "wide",
    "initial_sidebar_state": "collapsed",
}

_SCORE_TIERS: dict[str, dict[str, Any]] = {
    "strong": {"min": 70, "color": "#15C39A", "label": "Strong"},
    "medium": {"min": 50, "color": "#E8A735", "label": "Needs work"},
    "weak":   {"min": 0,  "color": "#E5534B", "label": "Needs improvement"},
}

_CUSTOM_CSS: str = """
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Instrument+Serif:ital@0;1&display=swap');

/* Base */
*, *::before, *::after { box-sizing: border-box; }
#MainMenu, footer, header { visibility: hidden; }
html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; color: #0D0F11; background: #F7F8FA; }
.main .block-container { padding: 1.5rem 2.5rem 3rem; max-width: 1440px; margin: 0 auto; }

/* the header & title */
.hero-wrap {
    background: #0a0a0f;
    border-radius: 0 0 28px 28px;
    padding: 5rem 3rem 5rem;
    margin: -1rem -3rem 2.5rem;
    position: relative;
    overflow: hidden;
}
.hero-wrap::before {
    content: '';
    position: absolute;
    top: -120px; right: -140px;
    width: 700px; height: 700px;
    background: radial-gradient(circle, rgba(139,92,246,0.18) 0%, transparent 70%);
    filter: blur(60px);
    pointer-events: none;
}
.hero-wrap::after {
    content: '';
    position: absolute;
    bottom: -120px; left: -40px;
    width: 700px; height: 700px;
    background: radial-gradient(circle, rgba(20,184,166,0.12) 0%, transparent 70%);
    filter: blur(60px);
    pointer-events: none;
}
.hero-eyebrow {
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 0.16em;
    text-transform: uppercase;
    color: #14b8a6;
    margin-bottom: 0.75rem;
}
.hero-title {
    color: #f8fafc !important;
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 72px !important;
    font-style: normal !important;
    line-height: 1 !important;
    letter-spacing: -0.02em !important;
    -webkit-font-smoothing: antialiased !important;
    text-rendering: optimizeLegibility !important;
}

.hero-title .title-prefix {
    font-style: normal;
    color: #f8fafc;
}
.hero-title .accent {
    font-style: italic;
    color: #a78bfa;
}
.hero-sub {
    font-size: 1.08rem;
    color: rgba(255,255,255,0.72);
    font-weight: 400;
    margin: 0;
    max-width: 640px;
}
.hero-badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: rgba(20,184,166,0.12);
    border: 1px solid rgba(20,184,166,0.25);
    color: #14b8a6;
    border-radius: 999px;
    padding: 0.3rem 0.85rem;
    font-size: 0.75rem;
    font-weight: 500;
    margin-top: 1.25rem;
}

/* UI Elements & Labels */
.section-label {
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 0.95rem !important;
    font-style: normal !important;
    text-transform: uppercase;
    letter-spacing: 0.02em !important;
    color: #6B7280; margin-bottom: 0.5rem; display: block;
    -webkit-font-smoothing: antialiased !important;
    text-rendering: optimizeLegibility !important;
}
.slim-divider {
    border: 0;
    height: 1px;
    margin: 0.9rem 0 1.15rem;
    background: linear-gradient(90deg, rgba(255,255,255,0), rgba(255,255,255,0.75), rgba(255,255,255,0));
    box-shadow: 0 1px 0 rgba(255,255,255,0.08);
}
.result-section-head {
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 0.95rem !important;
    font-style: normal !important;
    color: #0D0F11;
    padding: 0.75rem 0 0.4rem; display: flex; align-items: center; gap: 0.55rem;
    letter-spacing: 0.02em !important;
    -webkit-font-smoothing: antialiased !important;
    text-rendering: optimizeLegibility !important;
}
.result-section-head::after { content: ''; flex: 1; height: 1px; background: #E8EAED; }

/* Cards */
.card, .score-card {
    background: #FFFFFF; border: 1px solid #E8EAED; border-radius: 16px;
    padding: 1.25rem; margin-bottom: 1rem; box-shadow: 0 1px 3px rgba(13,15,17,0.04);
    transition: box-shadow 0.2s ease;
}
.card:hover { box-shadow: 0 4px 12px rgba(13,15,17,0.07); }

.setup-band {
    border: 0;
    background: transparent;
    padding: 0;
    margin: 0 0 1rem;
    box-shadow: none;
}

/* Score */
.score-ring-wrap { display: flex; align-items: center; gap: 1rem; padding: 0.25rem 0; }
.score-number {
    font-family: 'Instrument Serif', serif !important;
    font-size: 4rem;
    line-height: 1;
    font-weight: 400;
    letter-spacing: -0.03em;
}
.score-label-text {
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 0.95rem !important;
    font-style: normal !important;
    color: #0D0F11;
    letter-spacing: 0.02em !important;
    -webkit-font-smoothing: antialiased !important;
    text-rendering: optimizeLegibility !important;
}
.score-sub, .muted { font-size: 0.8rem; color: #6B7280; }

/* ── Contact & Status Chips ── */
.contact-grid, .kw-wrap, .review-toolbar { display: flex; flex-wrap: wrap; gap: 0.5rem; align-items: center; }
.contact-chip, .status-chip {
    background: #F7F8FA; border: 1px solid #E8EAED; border-radius: 999px;
    padding: 0.28rem 0.7rem; font-size: 0.76rem; color: #0D0F11;
    transition: background 0.15s;
}
.contact-chip:hover, .status-chip:hover { background: #E8EAED; }
.contact-chip b { color: #6B7280; font-weight: 700; margin-right: 0.25rem; }

/* PDF Viewer */
.pdf-shell {
    background: #2D3136; border: 1px solid #3C4148; border-radius: 14px;
    overflow: hidden; min-height: 720px; box-shadow: 0 4px 20px rgba(13,15,17,0.15);
}
.pdf-frame { width: 100%; min-height: 720px; border: 0; display: block; background: #2D3136; }

/* Suggestion Cards */
.suggestion-scroll { max-height: 760px; overflow-y: auto; padding-right: 0.3rem; }
.suggestion-scroll::-webkit-scrollbar { width: 5px; }
.suggestion-scroll::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 999px; }
.suggestion-card {
    background: #FFFFFF; border: 1px solid #E8EAED; border-radius: 12px;
    padding: 0; overflow: hidden; margin-bottom: 0.85rem;
    box-shadow: 0 1px 3px rgba(13,15,17,0.04); transition: all 0.2s ease;
}
.suggestion-card:hover { box-shadow: 0 4px 16px rgba(13,15,17,0.08); transform: translateY(-1px); }
.suggestion-card.accepted { border-color: #15C39A; background: #E6F8F3; }
.suggestion-card.dismissed { border-color: #fca5a5; background: #fef2f2; }
.suggestion-head { padding: 0.8rem 1rem 0; display: flex; justify-content: space-between; gap: 0.75rem; align-items: center; }
.suggestion-title { font-size: 0.72rem; font-weight: 700; color: #6B7280; text-transform: uppercase; letter-spacing: 0.08em; }
.fw-badge {
    background: rgba(21,195,154,0.1); color: #15C39A;
    font-size: 0.67rem; font-weight: 700; padding: 0.2rem 0.55rem;
    border-radius: 999px; white-space: nowrap; border: 1px solid rgba(21,195,154,0.3);
}
.severity-dot {
    width: 8px; height: 8px; border-radius: 50%; display: inline-block;
    margin-right: 0.35rem; vertical-align: middle; flex-shrink: 0;
}
.severity-dot.red { background: #E5534B; }
.severity-dot.yellow { background: #E8A735; }
.severity-dot.green { background: #15C39A; }

/* Rewrite segment */
.rewrite-grid {
    display: grid; grid-template-columns: minmax(0, 1fr) minmax(0, 1fr);
    gap: 0; border-top: 1px solid #E8EAED; margin-top: 0.7rem;
}
.rewrite-pane { padding: 0.85rem 1rem; min-width: 0; }
.rewrite-pane.before { background: #FAFAFA; border-right: 1px solid #E8EAED; }
.rewrite-pane.after { background: #FFFFFF; }
.pane-label {
    display: block; font-size: 0.64rem; font-weight: 700;
    letter-spacing: 0.1em; text-transform: uppercase; color: #6B7280; margin-bottom: 0.35rem;
}
.rewrite-text { font-size: 0.84rem; line-height: 1.6; color: #0D0F11; margin: 0; overflow-wrap: anywhere; }
.reasoning-row {
    padding: 0.72rem 1rem; border-top: 1px solid #E8EAED;
    background: #F7F8FA; color: #6B7280; font-size: 0.78rem; line-height: 1.5;
    font-style: italic;
}
.decision-bar { padding: 0.5rem 1rem; border-top: 1px solid #E8EAED; font-size: 0.74rem; font-weight: 700; }
.decision-bar.accepted { background: #15C39A; color: #FFFFFF; }
.decision-bar.dismissed { background: #fee2e2; color: #991b1b; }

/* Spacing */
.suggestion-btn-row { margin-top: 0.75rem; }

/* Keywords */
.kw-missing {
    background: #fef2f2; border: 1px solid #fecaca; color: #b91c1c;
    border-radius: 999px; padding: 0.28rem 0.7rem; font-size: 0.76rem; font-weight: 500;
    transition: all 0.15s;
}
.kw-missing:hover { background: #fee2e2; transform: scale(1.03); }
.kw-present {
    background: #E6F8F3; border: 1px solid rgba(21,195,154,0.3); color: #12A884;
    border-radius: 999px; padding: 0.28rem 0.7rem; font-size: 0.76rem; font-weight: 500;
    transition: all 0.15s;
}
.kw-present:hover { background: rgba(21,195,154,0.2); transform: scale(1.03); }

/* Alerts */
.warning-strip {
    background: #fffbeb; border-left: 3px solid #E8A735;
    border-radius: 0 10px 10px 0; padding: 0.65rem 0.9rem;
    font-size: 0.82rem; color: #78350f; margin-bottom: 0.5rem;
}
.ocr-strip {
    background: #eff6ff; border-left: 3px solid #3b82f6;
    border-radius: 0 10px 10px 0; padding: 0.65rem 0.9rem;
    font-size: 0.82rem; color: #1e40af; margin-bottom: 0.8rem;
}

/* stlit dashboard overrides */
[data-testid="stFileUploader"] {
    border: 2px dashed #E8EAED !important; border-radius: 12px !important;
    background: #F7F8FA !important; transition: border-color 0.2s !important;
}
[data-testid="stFileUploader"]:hover { border-color: #15C39A !important; }

/* Buttons */
.stButton > button,
.stButton > button * {
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.01em !important;
    text-rendering: optimizeLegibility !important;
    -webkit-font-smoothing: antialiased !important;
}
.stButton > button {
    background: linear-gradient(135deg, #8b5cf6, #7c3aed) !important;
    color: #f8f8f2 !important;
    border: 1px solid rgba(139, 92, 246, 0.35) !important;
    border-radius: 12px !important;
    white-space: nowrap !important;
    overflow: visible !important;
    text-overflow: clip !important;
    text-shadow: none !important;
    display: inline-block !important;
    box-shadow: 0 6px 18px rgba(124, 58, 237, 0.22) !important;
    transition: all 0.18s ease !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #7c3aed, #6d28d9) !important;
    border-color: rgba(139, 92, 246, 0.55) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 8px 22px rgba(124, 58, 237, 0.28) !important;
}
.stButton > button:active { transform: translateY(0) !important; }

.stDownloadButton > button {
    border-radius: 10px !important;
    background: #0D0F11 !important;
    color: #FFFFFF !important; border: none !important;
    box-shadow: 0 2px 8px rgba(13,15,17,0.2) !important;
    font-weight: 700 !important; transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    box-shadow: 0 4px 16px rgba(13,15,17,0.3) !important;
    transform: translateY(-1px) !important;
}

/* Analytics text */
[data-testid="stMetricLabel"],
[data-testid="stMetricValue"] {
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 48px !important;
    font-style: normal !important;
    letter-spacing: 0.02em !important;
    -webkit-font-smoothing: antialiased !important;
    text-rendering: optimizeLegibility !important;
}

/* Sidebar */
[data-testid="stSidebar"] { background-color: #1A1C1E; border-right: 1px solid #2A2D31; }
[data-testid="stSidebar"] p, [data-testid="stSidebar"] span, [data-testid="stSidebar"] label { color: #F7F8FA !important; }
[data-testid="stSidebar"] .card { background-color: #2D3136; border-color: #3C4148; }

textarea, input {
    border-radius: 10px !important; font-family: 'DM Sans', sans-serif !important;
    font-size: 0.84rem !important; border: 1px solid #E8EAED !important;
    transition: border-color 0.2s !important;
}
textarea:focus, input:focus { border-color: #15C39A !important; }
.stSpinner > div { border-top-color: #15C39A !important; }

/* Section completion cards */
.section-card {
    padding: 1rem; border-radius: 12px; text-align: center; color: white;
    font-weight: 700; transition: transform 0.2s; border: 1px solid transparent;
}
.section-card:hover { transform: scale(1.03); }
.section-card.present { background: #15C39A; border-color: #12A884; }
.section-card.missing { background: #E5534B; border-color: #C94A42; }

/* Navigation buttons */
div[role="radiogroup"] {
    background: linear-gradient(135deg, #0a0a0f 0%, #111827 45%, #172554 100%);
    padding: 0.45rem 0.55rem;
    border-radius: 16px;
    border: 1px solid rgba(139,92,246,0.25);
    gap: 0.4rem;
    margin: 0 auto 1.25rem;
    width: min(100%, 1180px);
    display: flex;
    justify-content: center;
    align-items: center;
    box-shadow: 0 10px 24px rgba(10,10,15,0.18);
}
label[data-baseweb="radio"] {
    background: rgba(255,255,255,0.06);
    padding: 0.55rem 0.85rem;
    border-radius: 10px;
    transition: all 0.18s ease;
    flex: 1 1 auto;
    text-align: center;
    min-width: 110px;
    cursor: pointer;
    border: 1px solid transparent;
    display: flex;
    align-items: center;
    justify-content: center;
}
label[data-baseweb="radio"]:hover {
    background: rgba(255,255,255,0.12);
    border-color: rgba(20,184,166,0.18);
}
label[data-baseweb="radio"] div:first-child { display: none; }
label[data-baseweb="radio"] *,
label[data-baseweb="radio"] > div,
label[data-baseweb="radio"] span {
    color: #ffffff !important;
    opacity: 1 !important;
    font-family: 'Instrument Serif', serif !important;
    font-weight: 400 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.02em !important;
    white-space: nowrap !important;
    overflow: visible !important;
    text-overflow: clip !important;
    text-shadow: none !important;
    display: inline-block !important;
}
label[data-baseweb="radio"][aria-checked="true"] {
    background: linear-gradient(135deg, rgba(139,92,246,0.22), rgba(20,184,166,0.14));
    border-color: rgba(139,92,246,0.45);
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.08);
}
label[data-baseweb="radio"][aria-checked="true"] div:last-child {
    color: #ffffff !important;
    text-shadow: 0 1px 3px rgba(0,0,0,0.35);
}

/* Score tips */
.score-tooltip-wrap {
    position: relative;
    cursor: pointer;
}
.score-tooltip-wrap .score-tooltip {
    visibility: hidden;
    opacity: 0;
    position: absolute;
    bottom: calc(100% + 10px);
    left: 50%;
    transform: translateX(-50%);
    background: #1A1C1E;
    color: #F7F8FA;
    border: 1px solid #3C4148;
    border-radius: 12px;
    padding: 0.75rem 1rem;
    font-size: 0.76rem;
    font-family: 'DM Sans', sans-serif;
    white-space: pre;
    line-height: 1.7;
    z-index: 999;
    box-shadow: 0 8px 24px rgba(0,0,0,0.35);
    pointer-events: none;
    transition: opacity 0.18s ease, visibility 0.18s ease;
    min-width: 200px;
}
.score-tooltip-wrap:hover .score-tooltip {
    visibility: visible;
    opacity: 1;
}
.score-tooltip-wrap .score-tooltip::after {
    content: '';
    position: absolute;
    top: 100%;
    left: 50%;
    transform: translateX(-50%);
    border: 6px solid transparent;
    border-top-color: #1A1C1E;
}

/* Model Selection menu */
.model-bar {
    display: flex;
    align-items: center;
    gap: 1.5rem;
    padding: 0.5rem 3rem;
    margin: -0.5rem -3rem 0;
    background: #0a0a0f;
    border-bottom: 1px solid rgba(255,255,255,0.06);
}
.model-bar .stSelectbox label,
.model-bar .stToggle label {
    color: rgba(255,255,255,0.7) !important;
    font-size: 0.78rem !important;
}

/* dynsizing */
@media (max-width: 900px) {
    .main .block-container { padding: 1rem; }
    .hero-wrap { border-radius: 0 0 18px 18px; padding: 2rem 1rem 1.25rem; margin: -1rem -1rem 1.25rem; }
    .hero-title { font-size: 2.4rem; }
    .hero-badge { display: inline-flex; margin-top: 0.9rem; }
    .rewrite-grid { grid-template-columns: 1fr; }
    .rewrite-pane.before { border-right: 0; border-bottom: 1px solid #E8EAED; }
    .pdf-frame, .pdf-shell { min-height: 560px; }
    .model-bar { padding: 0.5rem 1rem; margin: -0.5rem -1rem 0; }
}
"""

_SCORE_HISTORY_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "data",
    "score_history.json",
)


def _load_history() -> list[dict]:
    """load score history from disk."""
    try:
        if os.path.exists(_SCORE_HISTORY_PATH):
            with open(_SCORE_HISTORY_PATH, "r") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data[-50:]  # cap at 50
    except Exception:
        pass
    return []


def _save_history(history: list[dict]) -> None:
    """persist score history to disk."""
    try:
        os.makedirs(os.path.dirname(_SCORE_HISTORY_PATH), exist_ok=True)
        with open(_SCORE_HISTORY_PATH, "w") as f:
            json.dump(history[-50:], f)
    except Exception as e:
        print(f"failed to save score history: {e}")


# provider & model selection
_PROVIDER_MODELS: dict[str, list[str]] = {
    "Gemini": ["gemma-4-31b-it", "gemini-3.5-flash", "gemini-3.1-pro"],
    "Claude": ["claude-4-5-sonnet-latest", "claude-4-5-haiku-latest"],
    "ChatGPT": ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini"],
    "Local": ["llama3"],
}

_DISPLAY_TO_PROVIDER: dict[str, str] = {
    "Gemini": "gemini",
    "Claude": "claude",
    "ChatGPT": "chatgpt",
    "Local": "local",
}


def _model_opts() -> list[str]:
    """list of 'Provider -> model' for the dropdown."""
    opts = []
    for prov, models in _PROVIDER_MODELS.items():
        for m in models:
            opts.append(f"{prov} → {m}")
    return opts


def _parse_model_sel(sel: str) -> tuple[str, str, str]:
    """parse 'Provider → model' into (display, provider_key, model)."""
    parts = sel.split(" → ", 1)
    display = parts[0].strip()
    model = parts[1].strip() if len(parts) > 1 else ""
    return display, _DISPLAY_TO_PROVIDER.get(display, "gemini"), model


@dataclass
class AppState:
    results:  dict[str, Any] | None = None
    accepted: dict[str, bool] = field(default_factory=dict)
    analysed: bool = False
    score_history: list[dict] = field(default_factory=list)
    generated_cv: str = ""
    generated_cover_letter: str = ""
    pdf_bytes: bytes | None = None
    last_file_hash: int | None = None
    trigger_sidebar: bool = False
    active_suggestion_index: int = 0

def _init_state() -> None:
    """init streamlit session defaults."""
    defaults = AppState()
    for key, val in defaults.__dict__.items():
        if key not in st.session_state:
            st.session_state[key] = val
    # load persistent history on first init
    if not st.session_state.score_history:
        st.session_state.score_history = _load_history()

def _get_score_cfg(score: int) -> dict[str, Any]:
    """return tier config for a given score."""
    if score >= _SCORE_TIERS["strong"]["min"]:
        return _SCORE_TIERS["strong"]
    if score >= _SCORE_TIERS["medium"]["min"]:
        return _SCORE_TIERS["medium"]
    return _SCORE_TIERS["weak"]

def _hero() -> None:
    """render the top banner."""
    st.markdown("""
    <div class="hero-wrap">
        <div class="hero-eyebrow">Resume intelligence, reimagined</div>
        <h1 class="hero-title"><span class="title-prefix">RagsToRiches:</span><br><span class="accent">Smarter resumes, smarter opportunities.</span></h1>
        <p class="hero-sub">
            Review rewrite suggestions against the uploaded PDF, apply the changes you trust,
            then generate a CV from those decisions with or without a job description.
        </p>
    </div>
    """, unsafe_allow_html=True)

def _iter_rewrites(rewrites: dict[str, list[dict]]) -> list[tuple[str, int, dict]]:
    return [
        (section_name, i, item)
        for section_name, bullets in rewrites.items()
        for i, item in enumerate(bullets)
    ]

def _sug_key(sec: str, i: int, item: dict) -> str:
    return item.get("id") or f"{sec}_{i}"

def _actionable(rewrites: dict[str, list[dict]]) -> list[tuple[str, int, dict]]:
    return [
        (sec, i, item)
        for sec, i, item in _iter_rewrites(rewrites)
        if item.get("framework_used") not in ("none", "error")
        and item.get("original") != item.get("rewritten")
    ]

def _accepted_map(rewrites: dict[str, list[dict]]) -> dict[str, str]:
    amap: dict[str, str] = {}
    for sec, i, item in _iter_rewrites(rewrites):
        if not isinstance(item, dict):
            continue
        key = _sug_key(sec, i, item)
        if st.session_state.accepted.get(key) is True:
            amap[item.get("original", "")] = item.get("rewritten", "")
    return amap

def _generate_docx(md_text: str) -> bytes:
    """convert markdown cv text into a formatted docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re as _re
    except ImportError:
        raise RuntimeError("python-docx not installed. run: pip install python-docx")

    if not md_text:
        raise ValueError("generated cv is empty.")
    if len(md_text.strip()) < 100:
        raise ValueError("generated cv appears incomplete.")

    doc = Document()

    # tight margins to keep resume under 2 page 
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    # default style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    def _add_hr(doc):
        """add a thin hr via paragraph border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def _add_runs(p, raw_line: str, base_size: float = 10.5, base_bold: bool = False):
        """parse inline **bold** and *italic* markers."""
        import re as rr
        parts = rr.split(r'(\*\*.*?\*\*|\*.*?\*|_.*?_)', raw_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = p.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith('*') and part.endswith('*') and len(part) > 2) or \
                 (part.startswith('_') and part.endswith('_') and len(part) > 2):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)
                run.bold = base_bold
            run.font.size = Pt(base_size)

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # h1: candidate name
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x0D, 0x0F, 0x11)
            i += 1
            continue

        # h2: section headings
        if line.startswith('## '):
            heading_text = line[3:].strip().upper()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x0D, 0x0F, 0x11)
            _add_hr(doc)
            i += 1
            continue

        # h3: role/project titles
        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            _add_runs(p, line[4:].strip(), base_size=10.5, base_bold=True)
            i += 1
            continue

        # horizontal rule
        if _re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', line.strip()):
            _add_hr(doc)
            i += 1
            continue

        # bullet point
        bullet_match = _re.match(r'^(\s*)[-*•]\s+(.*)', line)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)
            _add_runs(p, bullet_match.group(2).strip(), base_size=10.5)
            i += 1
            continue

        # regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        _add_runs(p, line.strip(), base_size=10.5)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def _gen_pdf(md_text: str) -> bytes | None:
    """convert markdown to pdf bytes."""
    import tempfile
    import os

    try:
        from markdown_pdf import MarkdownPdf
        from markdown_pdf import Section
    except ImportError:
        try:
            from markdown_pdf import MarkdownPdf
            from markdown_pdf.Section import Section
        except ImportError:
            return None

    pdf = MarkdownPdf(toc_level=0)
    pdf.add_section(Section(md_text))

    tmp_path = os.path.join(tempfile.gettempdir(), "ragstoriches_cv.pdf")
    pdf.save(tmp_path)

    with open(tmp_path, "rb") as f:
        data = f.read()

    try:
        os.remove(tmp_path)
    except OSError:
        pass

    return data

_SEV_URGENCY: dict[str, str] = {
    "red":    "Needs immediate rewrite",
    "yellow": "Could be stronger",
    "green":  "Minor polish",
}

def _sev_color(severity: str, active: bool = False) -> tuple[float, float, float]:
    if active:
        return (0.04, 0.45, 0.28)
    if severity == "red":
        return (0.86, 0.24, 0.18)
    if severity == "green":
        return (0.10, 0.58, 0.33)
    return (0.92, 0.62, 0.05)

def _search_frags(text: str) -> list[str]:
    cleaned = " ".join(text.split())
    words = cleaned.split()
    phrases = [cleaned]
    if len(words) > 18:
        phrases.append(" ".join(words[:18]))
    if len(words) > 10:
        phrases.append(" ".join(words[:10]))
    if len(words) > 10:
        phrases.append(" ".join(words[-10:]))
    return [phrase for phrase in phrases if len(phrase) >= 20]

@st.cache_data(show_spinner=False)
def _highlight_pdf(
    pdf_bytes: bytes,
    rewrite_items: tuple[tuple[str, str, str, str, str, int], ...],
    active_key: str | None = None,
) -> tuple[bytes, int | None]:
    
    try:
        import fitz
    except Exception:
        return pdf_bytes, None

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return pdf_bytes, None

    active_page: int | None = None

    for item_id, text, severity, reasoning, rewritten, number in rewrite_items:
        is_active = (item_id == active_key)
        found = False
        for phrase in _search_frags(text):
            if found:
                break
            for page_index, page in enumerate(doc):
                matches = page.search_for(phrase, quads=True)
                if not matches:
                    continue
                color = _sev_color(severity, active=is_active)
                for quad in matches[:3]:
                    annot = page.add_highlight_annot(quad)
                    annot.set_colors(stroke=color)
                    urgency = _SEV_URGENCY.get(severity, "Suggestion")
                    if rewritten and rewritten != text:
                        popup_text = (
                            f"[#{number}] [{urgency}] Suggested rewrite:\n{rewritten}\n\n"
                            f"Why: {reasoning}"
                        )
                    elif reasoning:
                        popup_text = f"[#{number}] [{urgency}] {reasoning}"
                    else:
                        popup_text = f"[#{number}] Rewrite suggestion: {item_id}"
                    annot.set_info(content=popup_text, title="RAGsToRiches")
                    annot.update(opacity=0.55 if is_active else 0.28)
                found = True
                if is_active:
                    active_page = page_index + 1 
                break

    output = io.BytesIO()
    doc.save(output, garbage=4, deflate=True)
    doc.close()
    return output.getvalue(), active_page

def _pdf_viewer(pdf_bytes: bytes | None, rewrites: dict[str, list[dict]], active_key: str | None = None) -> None:
    if not pdf_bytes:
        st.markdown('<div class="card muted">Upload and analyse a PDF to see highlighted rewrite targets.</div>', unsafe_allow_html=True)
        return

    actionable_list = list(_actionable(rewrites))
    highlight_payload = tuple(
        (
            _sug_key(sec, i, item),
            item.get("highlight_text") or item.get("original", ""),
            item.get("severity", "yellow"),
            item.get("reasoning", ""),
            item.get("rewritten", ""),
            col_idx + 1,
        )
        for col_idx, (sec, i, item) in enumerate(actionable_list)
    )
    if highlight_payload:
        rendered, active_page = _highlight_pdf(pdf_bytes, highlight_payload, active_key)
    else:
        rendered, active_page = pdf_bytes, None

    encoded = base64.b64encode(rendered).decode("utf-8")
    page_fragment = f"&page={active_page}" if active_page else ""
    st.markdown(
        f"""
        <div class="pdf-shell">
            <iframe class="pdf-frame" src="data:application/pdf;base64,{encoded}#toolbar=1&navpanes=0{page_fragment}"></iframe>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if highlight_payload:
        st.markdown(
            '<span style="font-size:0.72rem;font-weight:700;color:#6B7280;text-transform:uppercase;'
            'letter-spacing:0.08em;display:block;margin-top:0.65rem;margin-bottom:0.35rem">'
            'Jump to highlight</span>',
            unsafe_allow_html=True,
        )
        chip_cols = st.columns(min(len(actionable_list), 8))
        for col_idx, (sec, i, item) in enumerate(actionable_list):
            item_key = _sug_key(sec, i, item)
            severity = item.get("severity", "yellow")
            urgency = _SEV_URGENCY.get(severity, "Suggestion")
            snippet = (item.get("highlight_text") or item.get("original", ""))[:60]
            is_active = item_key == active_key
            label = f"{'▶ ' if is_active else ''}{col_idx + 1}"
            with chip_cols[col_idx % min(len(actionable_list), 8)]:
                if st.button(
                    label,
                    key=f"jump_{item_key}",
                    use_container_width=True,
                    help=f"{urgency}: {snippet}",
                ):
                    active_item_idx = next(
                        (j for j, (s, ii, it) in enumerate(actionable_list) if _sug_key(s, ii, it) == item_key),
                        0,
                    )
                    st.session_state.active_suggestion_index = active_item_idx
                    st.rerun()

@st.fragment(run_every=5)
def _render_annotations(analysis_id: int, suggestion_key: str) -> None:
    anns = db.get_annotations(analysis_id)
    relevant = [a for a in anns if a["key"] == suggestion_key]
    if relevant:
        st.markdown('<span class="section-label" style="margin-top:1.5rem;">Discussion</span>', unsafe_allow_html=True)
        for a in relevant:
            st.markdown(
                f'<div style="background:#F7F8FA;padding:0.7rem;border-radius:10px;margin-bottom:0.5rem;font-size:0.85rem;border:1px solid #E8EAED;">'
                f'<span style="font-weight:700;color:#0D0F11;margin-right:0.5rem;">{html.escape(a["user"])}</span>'
                f'<span style="color:#6B7280;font-size:0.75rem;">{a["time"][:16].replace("T", " ")}</span><br>'
                f'<div style="margin-top:0.3rem;">{html.escape(a["comment"])}</div></div>',
                unsafe_allow_html=True,
            )

def _rewrites_tab(rewrites: dict[str, list[dict]], pdf_bytes: bytes | None) -> None:
    actionable = _actionable(rewrites)
    if not rewrites:
        st.markdown('<p class="muted">No experience or project sections were found.</p>', unsafe_allow_html=True)
        return

    acc_count = sum(
        1 for sec, i, item in actionable
        if st.session_state.accepted.get(_sug_key(sec, i, item)) is True
    )
    dis_count = sum(
        1 for sec, i, item in actionable
        if st.session_state.accepted.get(_sug_key(sec, i, item)) is False
    )

    st.markdown(
        f"""
        <div class="review-toolbar" style="margin-bottom:0.75rem">
            <span class="status-chip">{len(actionable)} suggested changes</span>
            <span class="status-chip">{acc_count} accepted</span>
            <span class="status-chip">{dis_count} dismissed</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    tool_a, tool_b, _ = st.columns([1, 1, 4])
    with tool_a:
        if st.button("Accept all", use_container_width=True):
            for sec, i, item in actionable:
                st.session_state.accepted[_sug_key(sec, i, item)] = True
            st.session_state.generated_cv = ""
            st.rerun()
    with tool_b:
        if st.button("Clear decisions", use_container_width=True):
            st.session_state.accepted = {}
            st.session_state.generated_cv = ""
            st.rerun()

    if not actionable:
        pdf_col, review_col = st.columns([1.18, 1], gap="large")
        with pdf_col:
            st.markdown('<span class="section-label">Highlighted PDF</span>', unsafe_allow_html=True)
            _pdf_viewer(pdf_bytes, rewrites, active_key=None)
        with review_col:
            st.markdown('<span class="section-label">Rewrite Decisions</span>', unsafe_allow_html=True)
            st.markdown(
                '<div class="card muted">No rewrite-worthy sentences were detected. Header/date lines were skipped.</div>',
                unsafe_allow_html=True,
            )
        return

    st.session_state.active_suggestion_index = max(
        0, min(st.session_state.active_suggestion_index, len(actionable) - 1)
    )
    idx = st.session_state.active_suggestion_index
    section_name, i, item = actionable[idx]
    key = _sug_key(section_name, i, item)

    pdf_col, review_col = st.columns([1.18, 1], gap="large")
    with pdf_col:
        st.markdown('<span class="section-label">Highlighted PDF</span>', unsafe_allow_html=True)
        _pdf_viewer(pdf_bytes, rewrites, active_key=key)

    with review_col:
        nav_left, nav_mid, nav_right = st.columns([1, 3, 1])
        with nav_left:
            if st.button("←", key="nav_prev", disabled=(idx == 0), use_container_width=True):
                st.session_state.active_suggestion_index -= 1
                st.rerun()
        with nav_mid:
            cur_state = st.session_state.accepted.get(key)
            cur_state_text = "✓ Accepted" if cur_state is True else "✗ Dismissed" if cur_state is False else "Undecided"
            st.markdown(
                f'<div style="text-align:center;padding:0.5rem 0">'
                f'<span class="status-chip">Suggestion {idx + 1} of {len(actionable)}</span>'
                f'&nbsp;&nbsp;<span class="status-chip">{html.escape(section_name.title())}</span>'
                f'&nbsp;&nbsp;<span class="status-chip">{html.escape(cur_state_text)}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
        with nav_right:
            if st.button("→", key="nav_next", disabled=(idx == len(actionable) - 1), use_container_width=True):
                st.session_state.active_suggestion_index += 1
                st.rerun()

        st.markdown('<span class="section-label">Rewrite Decision</span>', unsafe_allow_html=True)

        state = st.session_state.accepted.get(key)
        state_class = "accepted" if state is True else "dismissed" if state is False else ""
        state_text = "Accepted" if state is True else "Dismissed" if state is False else "Needs decision"
        safe_original = html.escape(item.get("original", ""))
        safe_rewritten = html.escape(item.get("rewritten", ""))
        safe_reasoning = html.escape(item.get("reasoning", ""))
        fw_badge = html.escape(item.get("framework_used", ""))

        severity = item.get("severity", "yellow")
        urgency_label = _SEV_URGENCY.get(severity, "Suggestion")
        sev_dot = f'<span class="severity-dot {severity}"></span>'

        st.markdown(
            f"""
            <div class="suggestion-card {state_class}">
                <div class="suggestion-head">
                    <span class="suggestion-title">#{idx + 1} &nbsp;{sev_dot}{html.escape(urgency_label)}</span>
                    <span class="fw-badge">{fw_badge}</span>
                </div>
                <div class="rewrite-grid">
                    <div class="rewrite-pane before">
                        <span class="pane-label">Original</span>
                        <p class="rewrite-text">{safe_original}</p>
                    </div>
                    <div class="rewrite-pane after">
                        <span class="pane-label">Suggested rewrite</span>
                        <p class="rewrite-text">{safe_rewritten}</p>
                    </div>
                </div>
                <div class="reasoning-row">💡 {safe_reasoning}</div>
            """,
            unsafe_allow_html=True,
        )

        if state is True:
            st.markdown('<div class="decision-bar accepted">This rewrite will be used in CV generation.</div></div>', unsafe_allow_html=True)
        elif state is False:
            st.markdown('<div class="decision-bar dismissed">Original text kept — suggested rewrite omitted from CV generation.</div></div>', unsafe_allow_html=True)
        else:
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="suggestion-btn-row">', unsafe_allow_html=True)
        btn_a, btn_b = st.columns(2)
        with btn_a:
            if st.button("Accept", key=f"acc_{key}", use_container_width=True):
                st.session_state.accepted[key] = True
                st.session_state.generated_cv = ""
                st.rerun()
        with btn_b:
            if st.button("Dismiss", key=f"rej_{key}", use_container_width=True):
                st.session_state.accepted[key] = False
                st.session_state.generated_cv = ""
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        if st.session_state.get("db_analysis_id"):
            _render_annotations(st.session_state.db_analysis_id, key)
            
            curr_user = st.session_state.get("current_user")
            ann_text = st.text_input("Add a comment", key=f"ann_in_{key}")
            if st.button("Post Comment", key=f"ann_btn_{key}"):
                if ann_text.strip() and curr_user:
                    db.save_annotation(st.session_state.db_analysis_id, curr_user["id"], key, ann_text)
                    st.rerun()

def _keywords_tab(jd_kws: list[str], missing: list[str], present: list[str], freqs: dict[str, int]) -> None:
    if not jd_kws:
        st.markdown(
            '<p style="color:#a1a1aa;padding:1rem 0">Paste a job description above and re-analyse to see keyword gaps.</p>',
            unsafe_allow_html=True,
        )
        return

    coverage = int(len(present) / len(jd_kws) * 100) if jd_kws else 0
    st.markdown(f"""
    <div class="card" style="margin-bottom:1.5rem">
        <span class="section-label">Coverage</span>
        <p style="font-size:2.2rem;font-weight:400;color:#0D0F11;margin:0;font-family:'Instrument Serif',serif;letter-spacing:0.02em;text-rendering:optimizeLegibility;-webkit-font-smoothing:antialiased;">
            {coverage}<span style="font-size:1rem;color:#6B7280;font-weight:400;font-family:'Instrument Serif',serif;">%</span>
        </p>
        <p style="font-size:0.82rem;color:#6B7280;margin:2px 0 0">
            {len(present)} of {len(jd_kws)} JD keywords present in your resume
        </p>
    </div>
    """, unsafe_allow_html=True)

    miss_col, have_col = st.columns(2, gap="large")

    with miss_col:
        st.markdown(f'<span class="section-label">❌ Missing ({len(missing)})</span>', unsafe_allow_html=True)
        chips = "".join(f'<span class="kw-missing">{html.escape(kw)}</span>' for kw in missing)
        st.markdown(
            f'<div class="kw-wrap">{chips or "<i style=\'color:#a1a1aa\'>None — great coverage!</i>"}</div>',
            unsafe_allow_html=True,
        )

    with have_col:
        st.markdown(f'<span class="section-label">✅ Present ({len(present)})</span>', unsafe_allow_html=True)
        chips = "".join(f'<span class="kw-present">{html.escape(kw)} <span style="opacity:0.7;font-size:0.8em">({freqs.get(kw, 0)})</span></span>' for kw in present)
        st.markdown(
            f'<div class="kw-wrap">{chips or "<i style=\'color:#a1a1aa\'>No matches found</i>"}</div>',
            unsafe_allow_html=True,
        )

def main() -> None:
    if not BACKEND_AVAILABLE:
        st.error("Missing backend logic. Ensure parser.py, analyser.py and router.py exist.")
        st.stop()

    st.set_page_config(**_PAGE_CONFIG)
    
    auth.render_auth_page()
    
    _init_state()
    if _CUSTOM_CSS.strip():
        st.markdown(f"<style>{_CUSTOM_CSS}</style>", unsafe_allow_html=True)

    _hero()

    model_options = _model_opts()
    bar_left, bar_right = st.columns([1, 1], gap="large")
    with bar_left:
        model_sel = st.selectbox(
            "LLM Model",
            options=model_options,
            index=0,
            help="Select provider and model.",
            label_visibility="collapsed",
        )
    with bar_right:
        use_critic = st.toggle(
            "Agentic Self-Correction",
            value=False,
            help="Runs a self-correction loop on rewritten bullets. Slower but higher quality.",
        )

    display_prov, sel_prov, sel_model = _parse_model_sel(model_sel)

    # providers env var names
    _KEY_NAMES = {"gemini": "GEMINI_API_KEY", "claude": "ANTHROPIC_API_KEY", "chatgpt": "OPENAI_API_KEY"}
    _ENV_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")

    local_endpoint = "http://localhost:11434/api/chat"
    if sel_prov == "local":
        local_endpoint = st.text_input(
            "Local API Endpoint",
            value="http://localhost:11434/api/chat",
            help="Default Ollama endpoint.",
        )
    elif sel_prov in _KEY_NAMES:
        env_var = _KEY_NAMES[sel_prov]
        api_key = os.environ.get(env_var)
        if not api_key or _is_key_placeholder(api_key):
            # in-app api key entry — writes to .env so it persists
            st.warning(f"No API key found for {display_prov}. Enter it below to get started.")
            api_key_input = st.text_input(
                f"{display_prov} API Key",
                type="password",
                placeholder=f"Paste your {display_prov} API key here",
                key=f"api_key_{sel_prov}",
            )
            if st.button(f"Save {display_prov} API Key", key=f"save_key_{sel_prov}"):
                if api_key_input.strip():
                    # read existing .env, update or append the key
                    lines = []
                    key_found = False
                    if os.path.exists(_ENV_PATH):
                        with open(_ENV_PATH, "r") as f:
                            lines = f.readlines()
                    new_lines = []
                    for line in lines:
                        if line.strip().startswith(f"{env_var}="):
                            new_lines.append(f"{env_var}={api_key_input.strip()}\n")
                            key_found = True
                        else:
                            new_lines.append(line)
                    if not key_found:
                        new_lines.append(f"\n{env_var}={api_key_input.strip()}\n")
                    with open(_ENV_PATH, "w") as f:
                        f.writelines(new_lines)
                    # hot-reload into current process
                    from dotenv import load_dotenv
                    load_dotenv(_ENV_PATH, override=True)
                    st.success(f"{display_prov} API key saved to .env")
                    st.rerun()
                else:
                    st.error("API key cannot be empty.")
            st.stop()

    st.markdown('<div class="setup-band">', unsafe_allow_html=True)
    st.markdown('<hr class="slim-divider">', unsafe_allow_html=True)
    left, right = st.columns([0.95, 1.25], gap="large")

    with left:
        st.markdown('<span class="section-label">Resume PDF / DOCX / TXT / MD</span>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "doc", "txt", "md"], label_visibility="collapsed", key="pdf_upload")
        if uploaded_file:
            st.success(f"{uploaded_file.name}")

    with right:
        st.markdown('<span class="section-label">Job Description <span style="color:#6B7280;font-weight:500;text-transform:none;letter-spacing:0">(optional for ATS matching)</span></span>', unsafe_allow_html=True)
        job_description = st.text_area(
            "Job Description",
            placeholder="Paste a full job description for keyword matching, or leave blank to improve the CV from rewrite decisions only.",
            height=152,
            label_visibility="collapsed",
        )
    st.markdown('</div>', unsafe_allow_html=True)

    if uploaded_file:
        uploaded_bytes = uploaded_file.getvalue()
        st.session_state.pdf_bytes = uploaded_bytes
        analyse_clicked = st.button("Analyse resume ✨", use_container_width=True)
        
        current_hash = hash(uploaded_bytes)
        stale = False
        if st.session_state.get("last_file_hash") != current_hash:
             st.session_state.results = None
             st.session_state.analysed = False
             st.session_state.last_file_hash = current_hash
             st.session_state.accepted = {}
             st.session_state.generated_cv = ""
             st.session_state.generated_cover_letter = ""

        if st.session_state.get("last_use_critic") is not None and st.session_state.get("last_use_critic") != use_critic:
             st.session_state.results = None
             st.session_state.analysed = False
             stale = True
             st.session_state.accepted = {}
             st.session_state.generated_cv = ""
             st.session_state.generated_cover_letter = ""
        st.session_state.last_use_critic = use_critic

        if stale and not analyse_clicked:
            st.info("Settings changed — click \"Analyse resume\" to re-run with the new settings.")

        if analyse_clicked:
            with st.spinner("Parsing resume..."):
                parsed = parse_file(uploaded_bytes, uploaded_file.name)
            if not parsed.sections and parsed.warnings:
                for w in parsed.warnings:
                    st.error(w)
                st.stop()

            with st.spinner(f"Running AI analysis via {display_prov} ({sel_model}) — this may take a few minutes..."):
                try:
                    results = analyse(
                        resume=parsed, 
                        job_description=job_description,
                        provider=sel_prov,
                        local_endpoint=local_endpoint,
                        use_critic=use_critic,
                        model=sel_model,
                    )
                    results["parsed_resume_obj"] = parsed
                    st.session_state.parsed_resume = parsed
                    st.session_state.pdf_bytes = uploaded_bytes
                    st.session_state.results  = results
                    # save score to persistent history
                    score_data = results["score"]
                    st.session_state.score_history.append(score_data)
                    if len(st.session_state.score_history) > 50:
                        st.session_state.score_history = st.session_state.score_history[-50:]
                    _save_history(st.session_state.score_history)

                    # DB persistence
                    curr_user = st.session_state.get("current_user")
                    if curr_user:
                        import json as _json
                        db_resume = db.save_resume(
                            user_id=curr_user["id"],
                            filename=uploaded_file.name,
                            raw_bytes=uploaded_bytes,
                            parsed_json=_json.dumps(parsed.sections)
                        )
                        score_total = score_data.get("total", 0) if isinstance(score_data, dict) else int(score_data)
                        db_analysis = db.save_analysis(
                            resume_id=db_resume.id,
                            user_id=curr_user["id"],
                            results_json=_json.dumps({"rewrites": results["rewrites"], "score": results["score"]}),
                            job_description=job_description,
                            provider=sel_prov,
                            model=sel_model,
                            score_total=score_total
                        )
                        st.session_state.db_analysis_id = db_analysis.id
                        st.session_state.db_resume_id = db_resume.id

                    st.session_state.accepted = {}
                    st.session_state.generated_cv = ""
                    st.session_state.generated_cover_letter = ""
                    st.session_state.analysed = True
                    st.session_state.trigger_sidebar = True
                    
                    if not feedback.is_feedback_silenced():
                        st.session_state.show_feedback = True

                except Exception as e:
                    st.error(f"Analysis failed. Did you configure {display_prov} correctly? Error: {e}")
                    st.stop()

            st.rerun()

    if st.session_state.get("show_feedback"):
        st.toast(f"Loving RAGsToRiches? [Tell us how we're doing]({feedback.get_forms_url()})", icon="⭐")
        feedback.silence_feedback()
        st.session_state.show_feedback = False
    else:
        st.markdown(
            '<p class="muted" style="text-align:center;padding:0.5rem 0">Upload a PDF resume above to get started.</p>',
            unsafe_allow_html=True,
        )

    if st.session_state.results:
        r = st.session_state.results

        st.markdown("<hr class='slim-divider'>", unsafe_allow_html=True)

        if r.get("ocr_used"):
            st.markdown(
                '<div class="ocr-strip">🔍 No text layer detected — OCR was used. '
                'Accuracy may be slightly lower on styled or image-heavy PDFs.</div>',
                unsafe_allow_html=True,
            )

        score_data = r.get("score", {"total": 0, "base": 30, "sections": 0, "keywords": 0, "bullet_quality": 0, "action_verbs": 0, "warnings": 0})
        if isinstance(score_data, int):
            score = score_data
            tooltip_html = "Score breakdown not available"
        else:
            score = score_data.get("total", 0)
            tooltip_lines = [
                f"Base:            {score_data.get('base', 0)}",
                f"Sections:      +{score_data.get('sections', 0)}",
                f"Keywords:      +{score_data.get('keywords', 0)}",
                f"Bullet Quality: +{score_data.get('bullet_quality', 0)}",
                f"Action Verbs:  +{score_data.get('action_verbs', 0)}",
                f"Warnings:       {score_data.get('warnings', 0)}",
                f"{'─' * 26}",
                f"Total:          {score}/100",
            ]
            tooltip_html = html.escape("\n".join(tooltip_lines))

        score_cfg = _get_score_cfg(score)

        st.markdown(f"""
        <div class="card" style="margin-bottom: 1rem; background: #FFFFFF; border: 1px solid #E8EAED; border-radius: 16px; padding: 1rem 1.1rem;">
            <span class="section-label">Resume Score</span>
            <div class="score-tooltip-wrap">
                <div class="score-ring-wrap" style="gap: 1rem; align-items: center;">
                    <div class="score-number" style="color:{score_cfg['color']};">{score}</div>
                    <div class="score-meta">
                        <span class="score-label-text">{score_cfg['label']}</span>
                        <span class="score-sub">out of 100 · hover for breakdown</span>
                    </div>
                </div>
                <div class="score-tooltip">{tooltip_html}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.sidebar:
            st.markdown(f"""
            <div class="card" style="margin-top:2rem">
                <span class="section-label">Resume Score</span>
                <div class="score-tooltip-wrap">
                    <div class="score-ring-wrap">
                        <div class="score-number" style="color:{score_cfg['color']}">{score}</div>
                        <div class="score-meta">
                            <span class="score-label-text">{score_cfg['label']}</span>
                            <span class="score-sub">out of 100</span>
                        </div>
                    </div>
                    <div class="score-tooltip">{tooltip_html}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            contact = r.get("contact", {})
            if contact:
                chips = "".join(
                    f'<span class="contact-chip"><b>{html.escape(k.title())}</b>{html.escape(str(v))}</span>'
                    for k, v in contact.items()
                )
                st.markdown(f"""
                <div class="card" style="margin-top:1rem">
                    <span class="section-label">Contact Detected</span>
                    <div class="contact-grid">{chips}</div>
                </div>
                """, unsafe_allow_html=True)
        
        
        warnings = [w for w in r.get("warnings", []) if "not detected" in w.lower() or "corrupt" in w.lower()]
        if warnings:
            st.markdown('<span class="section-label" style="margin-top:0.5rem">Parser Notes</span>', unsafe_allow_html=True)
            for w in warnings:
                st.markdown(f'<div class="warning-strip">{html.escape(w)}</div>', unsafe_allow_html=True)

        # parser debug sidebar
        with st.sidebar:
            st.markdown('<hr style="border-top:1px solid #3C4148">', unsafe_allow_html=True)
            st.markdown('<span class="section-label" style="color:#A1A5AB;">Parser Debug</span>', unsafe_allow_html=True)
            
            parsed_sections = r.get("sections", {})
            for sec in ["EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS"]:
                icon = "✓" if sec in parsed_sections else "✗"
                color = "#15C39A" if sec in parsed_sections else "#E5534B"
                st.markdown(f"<span style='color:{color}'>{icon} {sec.title()}</span>", unsafe_allow_html=True)
                
            with st.expander("Raw Parsed Output", expanded=False):
                st.json(parsed_sections)

            curr_user = st.session_state.get("current_user", {})
            if curr_user:
                st.markdown('<hr style="border-top:1px solid #3C4148">', unsafe_allow_html=True)
                st.markdown('<span class="section-label" style="color:#A1A5AB;">Collaborative Review</span>', unsafe_allow_html=True)
                if curr_user.get("role") == "candidate":
                    session_code = st.text_input("Enter Mentor Session Code")
                    if st.button("Join Session", key="join_sess_btn"):
                        if session_code:
                            rs = db.join_session(session_code, curr_user["id"])
                            if rs:
                                st.success(f"Joined session {session_code}!")
                                st.session_state.active_session_code = session_code
                            else:
                                st.error("Invalid or inactive session code.")
                elif curr_user.get("role") == "mentor":
                    if st.button("Export Session Report", key="export_mentor_btn"):
                        report_md = mentor.export_mentor_report(curr_user["id"])
                        st.download_button(
                            "Download Report (MD)",
                            data=report_md,
                            file_name="mentor_report.md",
                            mime="text/markdown"
                        )

        # navigation
        nav_options = ["Rewrite Suggestions", "Keyword Gap", "Extracted Sections", "Tailored CV", "Cover Letter", "Analytics", "Mentor Dashboard", "Job Matching"]
        
        curr_user = st.session_state.get("current_user", {})
        if curr_user.get("role") == "mentor":
            nav_options = ["Mentor Dashboard"] + [n for n in nav_options if n != "Mentor Dashboard"]
            
        selected_view = st.radio("", nav_options, horizontal=True, label_visibility="collapsed")

        if selected_view == "Rewrite Suggestions":
            _rewrites_tab(r.get("rewrites", {}), st.session_state.get("pdf_bytes"))

            # Save decisions to DB
            if st.session_state.get("db_analysis_id") and st.session_state.accepted:
                db.save_decisions(st.session_state.db_analysis_id, st.session_state.accepted)

        elif selected_view == "Keyword Gap":
            jd_kws  = r.get("jd_keywords", [])
            missing = r.get("missing_keywords", [])
            present = [kw for kw in jd_kws if kw not in missing]
            freqs = r.get("keyword_frequencies", {})
            _keywords_tab(jd_kws, missing, present, freqs)

        elif selected_view == "Extracted Sections":
            sections = r.get("sections", {})
            for section_name, lines in sections.items():
                with st.expander(section_name.title(), expanded=(section_name == "EXPERIENCE")):
                    st.code("\n".join(lines), language=None)

        elif selected_view == "Tailored CV":
            st.markdown('<span class="section-label">Generate Tailored CV</span>', unsafe_allow_html=True)
            st.markdown(
                '<p class="muted">The generated CV applies accepted rewrites, ignores dismissed rewrites, and uses the job description only when one was provided during analysis.</p>',
                unsafe_allow_html=True,
            )
            if st.button("Generate CV", key="gen_cv_btn"):
                acc_rw_map = _accepted_map(r.get("rewrites", {}))

                with st.spinner("Generating CV..."):
                    cv_text = gen_cv(
                        st.session_state.parsed_resume,
                        job_description,
                        acc_rw_map,
                        sel_prov,
                        local_endpoint,
                        rewrite_suggestions=r.get("rewrites", {}),
                        rewrite_decisions=st.session_state.accepted,
                        model=sel_model,
                    )
                    st.session_state.generated_cv = cv_text

            if st.session_state.get("generated_cv"):
                st.markdown("### ✏️ Edit Your CV")
                st.markdown(
                    '<p class="muted">Make any final edits below. Your changes will be reflected in the downloaded files.</p>',
                    unsafe_allow_html=True,
                )
                edited_cv = st.text_area(
                    "Edit CV",
                    value=st.session_state.generated_cv,
                    height=450,
                    label_visibility="collapsed",
                    key="cv_editor",
                )
                # sync edits back
                if edited_cv != st.session_state.generated_cv:
                    st.session_state.generated_cv = edited_cv

                st.markdown("### Preview")
                st.markdown(st.session_state.generated_cv)

                dl_col1, dl_col2, dl_col3 = st.columns(3)
                with dl_col1:
                    st.download_button(
                        "📄 Download Markdown",
                        st.session_state.generated_cv,
                        "tailored_cv.md",
                        key="dl_md",
                    )

                with dl_col2:
                    try:
                        docx_bytes = _generate_docx(st.session_state.generated_cv)
                        st.download_button(
                            "📝 Download DOCX",
                            docx_bytes,
                            "tailored_cv.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_docx",
                        )
                    except Exception as e:
                        st.error(f"DOCX export failed: {e}")

                with dl_col3:
                    try:
                        pdf_dl_bytes = _gen_pdf(st.session_state.generated_cv)
                        if pdf_dl_bytes:
                            st.download_button(
                                "📕 Download PDF",
                                pdf_dl_bytes,
                                "tailored_cv.pdf",
                                mime="application/pdf",
                                key="dl_pdf",
                            )
                    except Exception as e:
                        st.error(f"PDF export failed: {e}")

        elif selected_view == "Cover Letter":
            st.markdown('<span class="section-label">Generate Cover Letter</span>', unsafe_allow_html=True)
            st.markdown(
                '<p class="muted">Generate a professional cover letter tailored to the job description. '
                'A job description is recommended for best results but not required.</p>',
                unsafe_allow_html=True,
            )
            if st.button("Generate Cover Letter", key="gen_cl_btn"):
                with st.spinner("Generating cover letter..."):
                    cl_text = gen_cover_letter(
                        st.session_state.parsed_resume,
                        job_description,
                        sel_prov,
                        local_endpoint,
                        model=sel_model,
                    )
                    st.session_state.generated_cover_letter = cl_text

            if st.session_state.get("generated_cover_letter"):
                st.markdown("### ✏️ Edit Your Cover Letter")
                st.markdown(
                    '<p class="muted">Make any final edits below. Your changes will be reflected in the downloaded files.</p>',
                    unsafe_allow_html=True,
                )
                edited_cl = st.text_area(
                    "Edit Cover Letter",
                    value=st.session_state.generated_cover_letter,
                    height=400,
                    label_visibility="collapsed",
                    key="cl_editor",
                )
                if edited_cl != st.session_state.generated_cover_letter:
                    st.session_state.generated_cover_letter = edited_cl

                st.markdown("### Preview")
                st.markdown(st.session_state.generated_cover_letter)

                cl_col1, cl_col2, cl_col3 = st.columns(3)
                with cl_col1:
                    st.download_button(
                        "📄 Download Markdown",
                        st.session_state.generated_cover_letter,
                        "cover_letter.md",
                        key="dl_cl_md",
                    )

                with cl_col2:
                    try:
                        cl_docx = _generate_docx(st.session_state.generated_cover_letter)
                        st.download_button(
                            "📝 Download DOCX",
                            cl_docx,
                            "cover_letter.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_cl_docx",
                        )
                    except Exception as e:
                        st.error(f"DOCX export failed: {e}")

                with cl_col3:
                    try:
                        cl_pdf = _gen_pdf(st.session_state.generated_cover_letter)
                        if cl_pdf:
                            st.download_button(
                                "📕 Download PDF",
                                cl_pdf,
                                "cover_letter.pdf",
                                mime="application/pdf",
                                key="dl_cl_pdf",
                            )
                    except Exception as e:
                        st.error(f"PDF export failed: {e}")

        elif selected_view == "Analytics":
            st.markdown('<span class="section-label">User Analytics Dashboard</span>', unsafe_allow_html=True)
            if not r:
                st.info("Run your first analysis to see progress here.")
            else:
                st.markdown("### Resume Score History")
                history = st.session_state.score_history
                if len(history) > 1:
                    try:
                        import altair as alt
                        import pandas as pd

                        # extract totals
                        chart_data = []
                        for idx, entry in enumerate(history):
                            if isinstance(entry, dict):
                                chart_data.extend([
                                    {"Attempt": idx + 1, "Component": "Base", "Points": entry.get("base", 0)},
                                    {"Attempt": idx + 1, "Component": "Sections", "Points": entry.get("sections", 0)},
                                    {"Attempt": idx + 1, "Component": "Keywords", "Points": entry.get("keywords", 0)},
                                    {"Attempt": idx + 1, "Component": "Bullet Quality", "Points": entry.get("bullet_quality", 0)},
                                    {"Attempt": idx + 1, "Component": "Action Verbs", "Points": entry.get("action_verbs", 0)},
                                ])
                            else:
                                chart_data.append({"Attempt": idx + 1, "Component": "Total", "Points": int(entry)})

                        df = pd.DataFrame(chart_data)

                        chart = (
                            alt.Chart(df)
                            .mark_area(opacity=0.8)
                            .encode(
                                x=alt.X(
                                    "Attempt:Q",
                                    title="Attempt",
                                    scale=alt.Scale(domain=[1, max(len(history), 2)]),
                                    axis=alt.Axis(tickMinStep=1, format="d"),
                                ),
                                y=alt.Y(
                                    "Points:Q",
                                    title="Score Contribution",
                                    scale=alt.Scale(domain=[0, 100]),
                                ),
                                color=alt.Color("Component:N", scale=alt.Scale(scheme="category10")),
                                tooltip=["Attempt:Q", "Component:N", "Points:Q"],
                            )
                            .properties(height=250)
                            .configure_axis(
                                labelFontSize=11,
                                titleFontSize=12,
                            )
                        )
                        st.altair_chart(chart, use_container_width=True)
                    except ImportError:
                        # fallback if altair not installed
                        scores = [e.get("total", 0) if isinstance(e, dict) else int(e) for e in history]
                        st.line_chart(scores, height=150)
                elif len(history) == 1:
                    entry = history[0]
                    single_score = entry.get("total", 0) if isinstance(entry, dict) else int(entry)
                    st.info(f"First analysis score: **{single_score}/100**. Run more analyses to see a trend line.")
                else:
                    st.info("Run multiple analyses to see a trend line.")

                if isinstance(score_data, dict) and "section_scores" in score_data:
                    st.markdown("### Readability & Quality Heatmap")
                    sec_scores = score_data.get("section_scores", {})
                    if sec_scores:
                        try:
                            import altair as alt
                            import pandas as pd
                            hm_data = []
                            for sec_name, data in sec_scores.items():
                                hm_data.append({"Section": sec_name.title(), "Quality": data.get("quality", 0)})
                            if hm_data:
                                hm_df = pd.DataFrame(hm_data)
                                hm_chart = alt.Chart(hm_df).mark_rect().encode(
                                    x=alt.X("Section:N", title=None, axis=alt.Axis(labelAngle=0)),
                                    color=alt.Color("Quality:Q", scale=alt.Scale(scheme="greens"), title="Quality Score"),
                                    tooltip=["Section:N", "Quality:Q"]
                                ).properties(height=150)
                                st.altair_chart(hm_chart, use_container_width=True)
                        except ImportError:
                            pass

                st.markdown("### Per-Run Metrics")
                
                jd_kws = r.get("jd_keywords", [])
                missing = r.get("missing_keywords", [])
                present = [kw for kw in jd_kws if kw not in missing]
                coverage = (len(present) / len(jd_kws) * 100) if jd_kws else 0
                
                rewrites = r.get("rewrites", {})
                improved_bullets = sum(
                    1 for section in rewrites.values()
                    for item in section
                    if isinstance(item, dict)
                    and item.get("framework_used") not in ("none", "error")
                    and item.get("original") != item.get("rewritten")
                )
                
                cols = st.columns(2)
                with cols[0]:
                    if r.get("no_jd_provided"):
                        st.metric("Keyword Coverage", "N/A", help="No Job Description provided.")
                    else:
                        st.metric("Keyword Coverage", f"{coverage:.0f}%")
                with cols[1]:
                    st.metric("Bullets Improved", improved_bullets)

                timing = r.get("timing")
                if timing:
                    st.markdown("### Performance Metrics")
                    tcols = st.columns(4)
                    tcols[0].metric("Keywords", f"{timing.get('keywords_ms', 0)}ms")
                    tcols[1].metric("Retrieval", f"{timing.get('retrieval_ms', 0)}ms")
                    tcols[2].metric("Rewriting", f"{timing.get('rewriting_ms', 0)}ms")
                    tcols[3].metric("Total", f"{timing.get('total_ms', 0)}ms")

                st.markdown("### Section Completion")
                st.markdown("This highlights which sections are present in your resume.")
                
                cols = st.columns(3)
                sections = ["EXPERIENCE", "PROJECTS", "EDUCATION"]
                for i, sec in enumerate(sections):
                    with cols[i]:
                        has_sec = sec in r.get("sections", {})
                        css_class = "present" if has_sec else "missing"
                        status = "✓ Present" if has_sec else "✗ Missing"
                        st.markdown(f"""
                        <div class="section-card {css_class}">
                            {sec}<br>
                            <span style="font-size: 0.8em;">{status}</span>
                        </div>
                        """, unsafe_allow_html=True)

        elif selected_view == "Mentor Dashboard":
            if curr_user.get("role") == "mentor":
                mentor.render_mentor_dashboard(curr_user["id"])
            else:
                st.warning("Mentor Dashboard is only available to mentor accounts.")

        elif selected_view == "Job Matching":
            st.markdown('<span class="section-label">Recruitment Integration</span>', unsafe_allow_html=True)
            job_url = st.text_input("Enter Job Description URL (Indeed, LinkedIn, etc.)")
            if st.button("Scrape Job Description", key="scrape_jd_btn"):
                if job_url:
                    with st.spinner("Scraping job description..."):
                        jd_scraped = job_scraper.scrape_jd(job_url)
                    st.text_area("Scraped Job Description", value=jd_scraped, height=200)
                else:
                    st.error("Please enter a valid URL.")

            st.markdown('<hr class="slim-divider">', unsafe_allow_html=True)
            
            st.markdown("### LinkedIn Profile Import")
            if st.button("Import from LinkedIn", key="linkedin_btn"):
                st.info("OAuth integration placeholder. Enter LinkedIn URL to scrape public profile instead.")
            
            li_url = st.text_input("LinkedIn Profile URL (Public)")
            if st.button("Scrape LinkedIn Profile", key="scrape_li_btn"):
                if li_url:
                    with st.spinner("Scraping LinkedIn profile..."):
                        li_data = job_scraper.scrape_linkedin_profile(li_url)
                        if "error" in li_data:
                            st.error(li_data["error"])
                        else:
                            st.json(li_data)
                else:
                    st.error("Please enter a LinkedIn Profile URL.")

    if st.session_state.get("trigger_sidebar"):
        import streamlit.components.v1 as components
        components.html(
            """
            <script>
                const expandBtn = window.parent.document.querySelector('[data-testid="collapsedControl"]');
                if (expandBtn) {
                    expandBtn.click();
                }
            </script>
            """,
            height=0,
            width=0,
        )
        st.session_state.trigger_sidebar = False

if __name__ == "__main__":
    main()